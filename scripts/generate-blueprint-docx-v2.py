#!/usr/bin/env python3
"""
MITHQAL v25.0 — Memory-efficient .docx generator for the FINAL blueprint.
Streams the markdown line-by-line and builds the docx incrementally to avoid
holding the entire document in memory at once.
"""
import os
import sys
import time
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
except ImportError:
    os.system(f"{sys.executable} -m pip install python-docx --quiet")
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml

INPUT = "/home/z/my-project/docs/blueprint/mithqal-v25-FINAL-blueprint.md"
OUTPUT = "/home/z/my-project/docs/blueprint/mithqal-v25-FINAL-blueprint.docx"

COLOR_PRIMARY = RGBColor(0x1A, 0x36, 0x5F)
COLOR_GOLD = RGBColor(0xC5, 0x9B, 0x1E)
COLOR_DARK = RGBColor(0x2D, 0x2D, 0x2D)
COLOR_GRAY = RGBColor(0x6B, 0x6B, 0x6B)
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def add_horizontal_rule(doc, color="C59B1E", size="12"):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="{size}" w:space="1" w:color="{color}"/></w:pBdr>'
    )
    pPr.append(pBdr)


def set_cell_shading(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_formatted_runs(paragraph, text):
    """Add runs with **bold** and *italic* and `code` formatting."""
    import re
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = COLOR_DARK
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def parse_markdown_table(lines):
    rows = []
    for line in lines:
        line = line.strip()
        if line.startswith("|") and line.endswith("|"):
            cells = [c.strip() for c in line[1:-1].split("|")]
            if cells and not all(set(c) <= set("-: ") for c in cells):
                rows.append(cells)
    return rows


def add_styled_table(doc, rows):
    if not rows:
        return
    max_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j >= max_cols:
                break
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            add_formatted_runs(p, cell_text)
            for run in p.runs:
                run.font.size = Pt(9)
            if i == 0:
                cell._tc.get_or_add_tcPr()
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1A365F"/>')
                cell._tc.get_or_add_tcPr().append(shading)
                for run in p.runs:
                    run.font.color.rgb = COLOR_WHITE
                    run.bold = True
    doc.add_paragraph()


def main():
    start_time = time.time()
    print(f"Reading: {INPUT}", flush=True)
    file_size = os.path.getsize(INPUT)
    print(f"Input size: {file_size / 1024 / 1024:.1f} MB", flush=True)

    doc = Document()

    # Page setup (A4)
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10.5)
    font.color.rgb = COLOR_DARK
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.3
    pf.space_after = Pt(4)

    # Heading styles
    for level, (size, color) in enumerate(
        [(20, COLOR_PRIMARY), (16, COLOR_PRIMARY), (13, COLOR_PRIMARY), (11, COLOR_DARK)], 1
    ):
        h_style = doc.styles[f'Heading {level}']
        h_style.font.name = 'Calibri'
        h_style.font.size = Pt(size)
        h_style.font.bold = True
        h_style.font.color.rgb = color
        h_style.paragraph_format.space_before = Pt(14)
        h_style.paragraph_format.space_after = Pt(6)
        h_style.paragraph_format.keep_with_next = True

    # Title style
    title_style = doc.styles['Title']
    title_style.font.name = 'Calibri'
    title_style.font.size = Pt(32)
    title_style.font.bold = True
    title_style.font.color.rgb = COLOR_PRIMARY

    # Read file in streaming mode to avoid loading everything into memory at once
    with open(INPUT, "r", encoding="utf-8") as f:
        lines = f.readlines()

    total_lines = len(lines)
    print(f"Total lines: {total_lines}", flush=True)

    i = 0
    in_code_block = False
    code_lines = []
    in_html_block = False
    progress_interval = max(1000, total_lines // 50)

    while i < total_lines:
        line = lines[i].rstrip("\n")
        stripped = line.strip()

        # Progress report
        if i % progress_interval == 0 and i > 0:
            elapsed = time.time() - start_time
            pct = (i / total_lines) * 100
            eta = (elapsed / i) * (total_lines - i) if i > 0 else 0
            print(f"  Progress: {i}/{total_lines} ({pct:.1f}%) elapsed={elapsed:.0f}s eta={eta:.0f}s", flush=True)

        # Code blocks
        if stripped.startswith("```"):
            if in_code_block:
                code_text = "\n".join(code_lines)
                if code_text.strip():
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(0.5)
                    run = p.add_run(code_text)
                    run.font.name = "Consolas"
                    run.font.size = Pt(8.5)
                    run.font.color.rgb = COLOR_GRAY
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
                    p._p.get_or_add_tcPr() if hasattr(p._p, 'get_or_add_tcPr') else None
                    pPr = p._p.get_or_add_pPr()
                    pPr.append(shading)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # HTML comments
        if stripped.startswith("<!--"):
            in_html_block = True
            if stripped.endswith("-->"):
                in_html_block = False
            i += 1
            continue
        if in_html_block:
            if stripped.endswith("-->"):
                in_html_block = False
            i += 1
            continue

        # Horizontal rules
        if stripped in ("---", "===") or (stripped and all(c == "-" for c in stripped) and len(stripped) > 3):
            add_horizontal_rule(doc)
            i += 1
            continue

        # Tables
        if stripped.startswith("|") and i + 1 < total_lines and "|---" in lines[i + 1]:
            table_lines = []
            while i < total_lines and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].rstrip("\n"))
                i += 1
            rows = parse_markdown_table(table_lines)
            add_styled_table(doc, rows)
            continue

        # Headings
        if stripped.startswith("#"):
            # Count heading level
            level = 0
            while level < len(stripped) and stripped[level] == "#":
                level += 1
            text = stripped[level:].strip()
            if level <= 4 and text:
                try:
                    heading = doc.add_heading(text, level=min(level, 4))
                except Exception:
                    p = doc.add_paragraph()
                    run = p.add_run(text)
                    run.bold = True
                    run.font.size = Pt(14 - level)
            else:
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(10)
            i += 1
            continue

        # Blockquotes
        if stripped.startswith(">"):
            quote_text = stripped[1:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.75)
            p.paragraph_format.right_indent = Cm(0.5)
            add_formatted_runs(p, quote_text)
            for run in p.runs:
                run.font.italic = True
                run.font.color.rgb = COLOR_PRIMARY
            # Left border for blockquote
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="18" w:space="8" w:color="C59B1E"/></w:pBdr>'
            )
            pPr.append(pBdr)
            i += 1
            continue

        # List items
        if stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_runs(p, text)
            i += 1
            continue

        # Numbered lists
        if stripped and stripped[0].isdigit() and "." in stripped[:4]:
            text = stripped[stripped.index(".") + 1:].strip()
            p = doc.add_paragraph(style="List Number")
            add_formatted_runs(p, text)
            i += 1
            continue

        # Empty lines
        if not stripped:
            i += 1
            continue

        # Regular paragraphs
        p = doc.add_paragraph()
        add_formatted_runs(p, stripped)
        i += 1

    # Finalize
    elapsed = time.time() - start_time
    print(f"Writing: {OUTPUT}", flush=True)
    doc.save(OUTPUT)
    final_size = os.path.getsize(OUTPUT)
    print(f"Done. Size: {final_size:,} bytes ({final_size / 1024 / 1024:.1f} MB)", flush=True)
    print(f"Total time: {elapsed:.1f}s", flush=True)


if __name__ == "__main__":
    main()
