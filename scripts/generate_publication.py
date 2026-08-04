#!/usr/bin/env python3
"""
Generate publication-ready versions of the MITHQAL Blueprint v19.

Outputs (all in /home/z/my-project/docs/blueprint/publication/):
  1. mithqal-blueprint-v19.md     — GitHub-ready Markdown
  2. mithqal-blueprint-v19.html   — Self-contained HTML with sticky TOC sidebar
  3. mithqal-blueprint-v19.pdf    — Professional PDF (ReportLab)
  4. mithqal-blueprint-v19.docx   — Editable Word document (python-docx)

All content is preserved verbatim from the source blueprint.
"""

import os
import re
import sys
import html
from datetime import datetime
from collections import Counter

# ---------------------------------------------------------------------------
# Paths and constants
# ---------------------------------------------------------------------------

BLUEPRINT_PATH = "/home/z/my-project/docs/blueprint/blueprint.txt"
OUT_DIR = "/home/z/my-project/docs/blueprint/publication"
MD_PATH = os.path.join(OUT_DIR, "mithqal-blueprint-v19.md")
HTML_PATH = os.path.join(OUT_DIR, "mithqal-blueprint-v19.html")
PDF_PATH = os.path.join(OUT_DIR, "mithqal-blueprint-v19.pdf")
DOCX_PATH = os.path.join(OUT_DIR, "mithqal-blueprint-v19.docx")

VERSION = "v19"
DATE_STR = "2026-07-19"
DOC_TITLE = "MITHQAL Blueprint"
DOC_SUBTITLE = "Constitutional Monetary Infrastructure Specification"

# ---------------------------------------------------------------------------
# Line classification patterns
# ---------------------------------------------------------------------------

PART_RE = re.compile(r"^Part \d+:", re.UNICODE)
ARTICLE_RE = re.compile(r"^Article [IVX]+:", re.UNICODE)
INVARIANT_RE = re.compile(r"^Invariant \d+:", re.UNICODE)
INTRO_RE = re.compile(r"^Introduction\s+[\u2014\-]", re.UNICODE)
BULLET_RE = re.compile(r"^        \u2022\s+(.*)$", re.UNICODE)  # 8 spaces, bullet, whitespace
END_MARKER_RE = re.compile(r"^\[END OF", re.UNICODE)
NUMBERED_RE = re.compile(r"^\d+\.\s+\S", re.UNICODE)
TITLE_RE = re.compile(
    r"^MITHQAL v\d+\s+[\u2014\-]+\s+FINAL CONSOLIDATED IMPLEMENTATION SPECIFICATION\s*$",
    re.UNICODE,
)
ASCII_ART_RE = re.compile(r"^[\s\u2500-\u257F\u2550-\u256C+=|/<>\-*$:.]+$")

# Known subheader phrases that do NOT end with a colon
KNOWN_SUBHEADERS = {
    "Summary Table",
    "Constitutional Interpretation",
    "Core Principles of the Policy Framework",
    "Mandatory Register Fields",
    "Constituent Definitions",
    "Implementation Notes",
    "Governance Notes",
    "Standardised Stress Scenarios",
    "DOCUMENT STRUCTURE COMPLETED",
    "KEY DELIVERABLES ACHIEVED",
    "NEXT STEPS",
    "THE COMPLETE INSTITUTIONAL FRAMEWORK",
    "THE MITHQAL CONSTITUTIONAL PRINCIPLES",
    "FINAL DECLARATION",
    "STATUS: DOCUMENT COMPLETE",
    "What MITHQAL Is",
    "What MITHQAL Is Not",
    "The Five Documentation Layers",
    "The Institution exists.",
}

# ---------------------------------------------------------------------------
# Parser
# ---------------------------------------------------------------------------

def classify_line(line, line_num):
    """Classify a single line into (type, content)."""
    stripped = line.rstrip("\n")
    if not stripped.strip():
        return ("blank", "")

    if END_MARKER_RE.match(stripped):
        return ("end_marker", stripped)

    m = BULLET_RE.match(line)
    if m:
        return ("bullet", m.group(1).rstrip())

    if PART_RE.match(stripped):
        return ("part", stripped)

    if ARTICLE_RE.match(stripped):
        return ("article", stripped)

    if INTRO_RE.match(stripped):
        return ("introduction", stripped)

    if INVARIANT_RE.match(stripped):
        return ("invariant", stripped)

    if TITLE_RE.match(stripped):
        if line_num == 1:
            return ("title", stripped)
        return ("section_divider", stripped)

    if NUMBERED_RE.match(stripped):
        return ("numbered", stripped)

    # ASCII art (box-drawing characters)
    if ASCII_ART_RE.match(stripped) and len(stripped) > 5 and any(
        ord(c) >= 0x2500 for c in stripped
    ):
        return ("ascii_art", stripped)

    # Subheader: short line ending with colon
    if len(stripped) < 100 and stripped.endswith(":") and " " in stripped:
        return ("subheader", stripped)

    # Known subheaders
    if stripped in KNOWN_SUBHEADERS:
        return ("subheader", stripped)

    # Short standalone label-like lines (no ending punctuation) that are likely
    # headers — only treat as subheader if they look like Title Case or are short
    # labels matching common patterns. Be conservative.
    if (
        len(stripped) < 60
        and not stripped.endswith((".", ",", ";", "?", "!"))
        and re.match(r"^[A-Z]", stripped)
        and "  " not in stripped
    ):
        return ("maybe_subheader", stripped)

    return ("paragraph", stripped)


def parse_blueprint(path):
    """Read blueprint and return list of (type, content) tuples."""
    with open(path, "r", encoding="utf-8") as f:
        raw_lines = f.readlines()

    parsed = []
    for i, line in enumerate(raw_lines, start=1):
        parsed.append(classify_line(line, i))

    # Post-process: convert maybe_subheader to subheader only if followed by
    # content (paragraph or bullet), otherwise convert to paragraph.
    n = len(parsed)
    result = []
    for i, (typ, content) in enumerate(parsed):
        if typ == "maybe_subheader":
            j = i + 1
            while j < n and parsed[j][0] == "blank":
                j += 1
            if j < n and parsed[j][0] in ("paragraph", "bullet", "numbered", "subheader"):
                result.append(("subheader", content))
            else:
                result.append(("paragraph", content))
        else:
            result.append((typ, content))

    return result


# ---------------------------------------------------------------------------
# Anchor / slug generation
# ---------------------------------------------------------------------------

SLUG_RE = re.compile(r"[^a-z0-9]+")


def slugify(text):
    text = text.lower()
    text = text.replace("\u2014", "-").replace("\u2013", "-")
    text = SLUG_RE.sub("-", text)
    return text.strip("-")


def make_part_anchor(part_text):
    m = re.match(r"^Part (\d+):", part_text)
    if m:
        return f"part-{m.group(1)}"
    return slugify(part_text)


def make_article_anchor(part_num, article_text):
    m = re.match(r"^Article ([IVX]+):", article_text)
    if m:
        return f"part-{part_num}-article-{m.group(1).lower()}"
    return slugify(article_text)


# ---------------------------------------------------------------------------
# Markdown generation
# ---------------------------------------------------------------------------

def collect_toc(parsed):
    """Return list of (level, title, anchor) — first occurrence of each Part."""
    toc = []
    seen_parts = set()
    current_part_num = None
    for typ, content in parsed:
        if typ == "part":
            m = re.match(r"^Part (\d+):", content)
            if m:
                pn = m.group(1)
                current_part_num = pn
                if pn not in seen_parts:
                    seen_parts.add(pn)
                    toc.append(("part", content, make_part_anchor(content)))
        elif typ == "article":
            if current_part_num:
                toc.append(("article", content, make_article_anchor(current_part_num, content)))
    return toc


def generate_markdown(parsed, md_path):
    toc = collect_toc(parsed)
    lines = []
    lines.append(f"# {DOC_TITLE} ({VERSION})")
    lines.append("")
    lines.append(f"**{DOC_SUBTITLE}**")
    lines.append("")
    lines.append(f"- **Version**: {VERSION} (v18 base + v19 constitutional evolution)")
    lines.append(f"- **Date**: {DATE_STR}")
    lines.append(f"- **Source**: `docs/blueprint/blueprint.txt` (28,456 lines)")
    lines.append(f"- **Structure**: 5 Parts (Layers), 49 Articles total")
    lines.append("")
    lines.append("---")
    lines.append("")
    lines.append("## Table of Contents")
    lines.append("")

    for level, title, anchor in toc:
        if level == "part":
            lines.append(f"- **[{title}](#{anchor})**")
        else:
            lines.append(f"  - [{title}](#{anchor})")

    lines.append("")
    lines.append("---")
    lines.append("")

    current_part_num = None
    in_ascii_block = False
    last_was_blank = True

    def emit(s):
        nonlocal last_was_blank
        if s == "":
            if last_was_blank:
                return
            last_was_blank = True
        else:
            last_was_blank = False
        lines.append(s)

    for typ, content in parsed:
        if typ == "title":
            continue
        if typ == "blank":
            emit("")
            continue
        if typ == "ascii_art":
            if not in_ascii_block:
                emit("")
                emit("```")
                in_ascii_block = True
            emit(content)
            continue
        else:
            if in_ascii_block:
                emit("```")
                emit("")
                in_ascii_block = False

        if typ == "part":
            m = re.match(r"^Part (\d+):", content)
            if m:
                current_part_num = m.group(1)
            anchor = make_part_anchor(content)
            emit("")
            emit(f'<a id="{anchor}"></a>')
            emit("")
            emit(f"# {content}")
            emit("")
        elif typ == "introduction":
            emit("")
            emit(f"## {content}")
            emit("")
        elif typ == "article":
            if current_part_num:
                anchor = make_article_anchor(current_part_num, content)
            else:
                anchor = slugify(content)
            emit("")
            emit(f'<a id="{anchor}"></a>')
            emit("")
            emit(f"## {content}")
            emit("")
        elif typ in ("invariant", "subheader", "numbered"):
            emit("")
            emit(f"### {content}")
            emit("")
        elif typ == "section_divider":
            emit("")
            emit(f"# {content}")
            emit("")
        elif typ == "end_marker":
            emit("")
            emit(f"_{content}_")
            emit("")
        elif typ == "bullet":
            emit(f"- {content}")
        elif typ == "paragraph":
            emit(content)

    if in_ascii_block:
        emit("```")
        emit("")

    with open(md_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines) + "\n")

    return len(lines)


# ---------------------------------------------------------------------------
# HTML generation
# ---------------------------------------------------------------------------

HTML_CSS = """
:root {
  --bg: #ffffff;
  --fg: #1a1a1a;
  --muted: #666;
  --accent: #1a5276;
  --accent-bg: #eaf2f8;
  --border: #d0d7de;
  --code-bg: #f6f8fa;
  --sidebar-bg: #fafbfc;
}
@media (prefers-color-scheme: dark) {
  :root {
    --bg: #0d1117;
    --fg: #e6edf3;
    --muted: #9198a1;
    --accent: #58a6ff;
    --accent-bg: #1a2332;
    --border: #30363d;
    --code-bg: #161b22;
    --sidebar-bg: #161b22;
  }
}
* { box-sizing: border-box; }
html, body {
  margin: 0;
  padding: 0;
  background: var(--bg);
  color: var(--fg);
  font-family: Georgia, "Times New Roman", serif;
  font-size: 16px;
  line-height: 1.65;
  -webkit-font-smoothing: antialiased;
}
.container { display: flex; min-height: 100vh; }
.sidebar {
  width: 320px;
  position: sticky;
  top: 0;
  height: 100vh;
  overflow-y: auto;
  background: var(--sidebar-bg);
  border-right: 1px solid var(--border);
  padding: 24px 20px;
  flex-shrink: 0;
  font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Helvetica, Arial, sans-serif;
  font-size: 14px;
}
.sidebar h2 {
  font-size: 16px;
  margin: 0 0 12px 0;
  color: var(--accent);
  font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Helvetica, Arial, sans-serif;
}
.sidebar details > summary {
  cursor: pointer;
  font-weight: 600;
  color: var(--accent);
  padding: 4px 0;
  list-style: none;
}
.sidebar details > summary::-webkit-details-marker { display: none; }
.sidebar details > summary::before {
  content: "\\25BC";
  display: inline-block;
  margin-right: 6px;
  font-size: 10px;
  transition: transform 0.2s;
}
.sidebar details[open] > summary::before { transform: rotate(-90deg); }
.sidebar ul { list-style: none; padding-left: 14px; margin: 4px 0 8px 0; }
.sidebar li { margin: 2px 0; }
.sidebar a {
  color: var(--fg);
  text-decoration: none;
  display: block;
  padding: 2px 6px;
  border-radius: 3px;
  font-size: 13px;
}
.sidebar a:hover { background: var(--accent-bg); color: var(--accent); }
.sidebar a.part-link { font-weight: 600; font-size: 14px; color: var(--accent); }
.content {
  flex: 1;
  padding: 60px 80px;
  max-width: 920px;
  margin: 0 auto;
  min-width: 0;
}
.doc-header {
  border-bottom: 2px solid var(--accent);
  padding-bottom: 24px;
  margin-bottom: 40px;
}
.doc-header h1 {
  font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Helvetica, Arial, sans-serif;
  font-size: 36px;
  margin: 0 0 8px 0;
  color: var(--accent);
  letter-spacing: -0.5px;
}
.doc-header .subtitle { font-size: 18px; color: var(--muted); margin: 0 0 12px 0; font-style: italic; }
.doc-header .meta {
  font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Helvetica, Arial, sans-serif;
  font-size: 13px;
  color: var(--muted);
}
.doc-header .meta span { margin-right: 16px; }
h1 {
  font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Helvetica, Arial, sans-serif;
  color: var(--accent);
  font-size: 28px;
  margin: 48px 0 16px 0;
  padding-bottom: 8px;
  border-bottom: 1px solid var(--border);
}
h2 {
  font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Helvetica, Arial, sans-serif;
  color: var(--accent);
  font-size: 22px;
  margin: 36px 0 14px 0;
}
h3 {
  font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Helvetica, Arial, sans-serif;
  color: var(--fg);
  font-size: 17px;
  font-weight: 600;
  margin: 24px 0 10px 0;
}
p { margin: 8px 0; }
ul { margin: 8px 0; padding-left: 24px; }
li { margin: 4px 0; }
pre {
  background: var(--code-bg);
  border: 1px solid var(--border);
  border-radius: 4px;
  padding: 12px 16px;
  overflow-x: auto;
  font-family: "SF Mono", Monaco, Consolas, "Courier New", monospace;
  font-size: 13px;
  line-height: 1.4;
  color: var(--fg);
}
code {
  font-family: "SF Mono", Monaco, Consolas, "Courier New", monospace;
  background: var(--code-bg);
  padding: 1px 5px;
  border-radius: 3px;
  font-size: 13px;
}
pre code { background: none; padding: 0; }
.end-marker {
  color: var(--muted);
  font-style: italic;
  font-size: 13px;
  text-align: center;
  margin: 20px 0;
  font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Helvetica, Arial, sans-serif;
}
.section-divider {
  color: var(--accent);
  font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Helvetica, Arial, sans-serif;
  font-size: 24px;
  text-align: center;
  margin: 40px 0;
  padding: 16px;
  border-top: 1px solid var(--border);
  border-bottom: 1px solid var(--border);
}
@media print {
  .sidebar { display: none; }
  .content { padding: 20px; max-width: 100%; }
  body { font-size: 11pt; line-height: 1.4; }
  h1 { page-break-before: always; }
  h1:first-of-type { page-break-before: avoid; }
  a { color: black; text-decoration: none; }
  pre { white-space: pre-wrap; word-wrap: break-word; }
}
@media (max-width: 1100px) {
  .sidebar { width: 260px; padding: 16px 12px; }
  .content { padding: 40px 32px; }
}
@media (max-width: 768px) {
  .container { flex-direction: column; }
  .sidebar { position: relative; width: 100%; height: auto; max-height: 300px; }
  .content { padding: 20px; }
}
"""


def esc(s):
    return html.escape(s, quote=True)


def generate_html(parsed, html_path):
    toc = collect_toc(parsed)
    # Group articles by part
    parts_with_articles = []
    current = None
    for level, title, anchor in toc:
        if level == "part":
            if current:
                parts_with_articles.append(current)
            current = {"title": title, "anchor": anchor, "articles": []}
        else:
            if current:
                current["articles"].append((title, anchor))
    if current:
        parts_with_articles.append(current)

    out = []
    out.append("<!DOCTYPE html>")
    out.append('<html lang="en">')
    out.append("<head>")
    out.append('<meta charset="UTF-8">')
    out.append('<meta name="viewport" content="width=device-width, initial-scale=1.0">')
    out.append(f"<title>{esc(DOC_TITLE)} {VERSION}</title>")
    out.append(f"<style>{HTML_CSS}</style>")
    out.append("</head>")
    out.append("<body>")
    out.append('<div class="container">')

    # Sidebar
    out.append('<aside class="sidebar">')
    out.append("<h2>Contents</h2>")
    for p in parts_with_articles:
        out.append("<details open>")
        out.append(f'<summary><a href="#{p["anchor"]}" class="part-link">{esc(p["title"])}</a></summary>')
        out.append("<ul>")
        for art_title, art_anchor in p["articles"]:
            out.append(f'<li><a href="#{art_anchor}">{esc(art_title)}</a></li>')
        out.append("</ul>")
        out.append("</details>")
    out.append("</aside>")

    # Main content
    out.append('<main class="content">')
    out.append('<header class="doc-header">')
    out.append(f"<h1>{esc(DOC_TITLE)} <span style='font-size:0.6em;color:var(--muted)'>{VERSION}</span></h1>")
    out.append(f'<p class="subtitle">{esc(DOC_SUBTITLE)}</p>')
    out.append('<div class="meta">')
    out.append(f"<span><strong>Version:</strong> {VERSION} (v18 base + v19 constitutional evolution)</span>")
    out.append(f"<span><strong>Date:</strong> {DATE_STR}</span>")
    out.append(f"<span><strong>Articles:</strong> 49 (5 Parts)</span>")
    out.append("</div>")
    out.append("</header>")

    current_part_num = None
    in_ascii = False
    in_ul = False

    def close_ul():
        nonlocal in_ul
        if in_ul:
            out.append("</ul>")
            in_ul = False

    def close_ascii():
        nonlocal in_ascii
        if in_ascii:
            out.append("</code></pre>")
            in_ascii = False

    for typ, content in parsed:
        if typ == "title":
            continue
        if typ == "blank":
            close_ul()
            close_ascii()
            continue
        if typ == "ascii_art":
            if not in_ascii:
                close_ul()
                out.append("<pre><code>")
                in_ascii = True
            out.append(esc(content))
            continue
        else:
            close_ascii()

        if typ == "part":
            close_ul()
            m = re.match(r"^Part (\d+):", content)
            if m:
                current_part_num = m.group(1)
            anchor = make_part_anchor(content)
            out.append(f'<h1 id="{anchor}">{esc(content)}</h1>')
        elif typ == "introduction":
            close_ul()
            out.append(f"<h2>{esc(content)}</h2>")
        elif typ == "article":
            close_ul()
            if current_part_num:
                anchor = make_article_anchor(current_part_num, content)
            else:
                anchor = slugify(content)
            out.append(f'<h2 id="{anchor}">{esc(content)}</h2>')
        elif typ in ("invariant", "subheader", "numbered"):
            close_ul()
            out.append(f"<h3>{esc(content)}</h3>")
        elif typ == "section_divider":
            close_ul()
            out.append(f'<div class="section-divider">{esc(content)}</div>')
        elif typ == "end_marker":
            close_ul()
            out.append(f'<div class="end-marker">{esc(content)}</div>')
        elif typ == "bullet":
            if not in_ul:
                out.append("<ul>")
                in_ul = True
            out.append(f"<li>{esc(content)}</li>")
        elif typ == "paragraph":
            close_ul()
            out.append(f"<p>{esc(content)}</p>")

    close_ul()
    close_ascii()

    out.append("</main>")
    out.append("</div>")
    out.append("</body>")
    out.append("</html>")

    with open(html_path, "w", encoding="utf-8") as f:
        f.write("\n".join(out))

    return len(out)


# ---------------------------------------------------------------------------
# DOCX generation
# ---------------------------------------------------------------------------

def generate_docx(parsed, docx_path):
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Configure Normal style
    normal = doc.styles["Normal"]
    normal.font.name = "Georgia"
    normal.font.size = Pt(11)

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(20)
    h1.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(16)
    h2.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)

    h3 = doc.styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(13)
    h3.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Title page
    for _ in range(6):
        doc.add_paragraph("")
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(DOC_TITLE)
    run.font.name = "Calibri"
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)
    run.bold = True

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_p.add_run(VERSION)
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    sub2_p = doc.add_paragraph()
    sub2_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub2_p.add_run(DOC_SUBTITLE)
    run.font.name = "Georgia"
    run.font.size = Pt(14)
    run.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    for _ in range(2):
        doc.add_paragraph("")

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta_p.add_run(
        f"Version: {VERSION} (v18 base + v19 constitutional evolution)\n"
        f"Date: {DATE_STR}\n"
        f"5 Parts (Layers) \u00b7 49 Articles"
    )
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # Table of Contents (manual listing)
    toc_heading = doc.add_paragraph()
    run = toc_heading.add_run("Table of Contents")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)

    toc = collect_toc(parsed)
    for level, title, anchor in toc:
        if level == "part":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            run = p.add_run(title)
            run.bold = True
            run.font.name = "Calibri"
            run.font.size = Pt(12)
        else:
            p = doc.add_paragraph(title)
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_after = Pt(2)
            for run in p.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(11)

    doc.add_page_break()

    # Body
    for typ, content in parsed:
        if typ == "title":
            continue
        if typ == "blank":
            continue
        if typ == "ascii_art":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            run = p.add_run(content)
            run.font.name = "Consolas"
            run.font.size = Pt(8)
            continue
        if typ == "part":
            doc.add_heading(content, level=1)
        elif typ in ("introduction", "article"):
            doc.add_heading(content, level=2)
        elif typ in ("invariant", "subheader", "numbered"):
            doc.add_heading(content, level=3)
        elif typ == "section_divider":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(content)
            run.bold = True
            run.font.name = "Calibri"
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)
        elif typ == "end_marker":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(content)
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        elif typ == "bullet":
            doc.add_paragraph(content, style="List Bullet")
        elif typ == "paragraph":
            doc.add_paragraph(content)

    doc.save(docx_path)


# ---------------------------------------------------------------------------
# PDF generation (ReportLab)
# ---------------------------------------------------------------------------

def generate_pdf(parsed, pdf_path):
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.colors import HexColor
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
    from reportlab.platypus import (
        BaseDocTemplate,
        PageTemplate,
        Frame,
        Paragraph,
        Spacer,
        PageBreak,
    )
    from reportlab.platypus.tableofcontents import TableOfContents

    PAGE_W, PAGE_H = letter

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "CoverTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=36,
        leading=44,
        alignment=TA_CENTER,
        textColor=HexColor("#1A5276"),
        spaceAfter=12,
    )

    subtitle_style = ParagraphStyle(
        "CoverSubtitle",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=18,
        leading=24,
        alignment=TA_CENTER,
        textColor=HexColor("#666666"),
        spaceAfter=8,
    )

    cover_meta_style = ParagraphStyle(
        "CoverMeta",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=11,
        leading=16,
        alignment=TA_CENTER,
        textColor=HexColor("#666666"),
    )

    h1_style = ParagraphStyle(
        "H1",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=20,
        leading=26,
        textColor=HexColor("#1A5276"),
        spaceBefore=18,
        spaceAfter=12,
        keepWithNext=True,
    )

    h2_style = ParagraphStyle(
        "H2",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=15,
        leading=20,
        textColor=HexColor("#1A5276"),
        spaceBefore=14,
        spaceAfter=8,
        keepWithNext=True,
    )

    h3_style = ParagraphStyle(
        "H3",
        parent=styles["Heading3"],
        fontName="Helvetica-Bold",
        fontSize=12,
        leading=16,
        textColor=HexColor("#333333"),
        spaceBefore=10,
        spaceAfter=6,
        keepWithNext=True,
    )

    body_style = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontName="Times-Roman",
        fontSize=10.5,
        leading=14,
        alignment=TA_JUSTIFY,
        spaceAfter=4,
    )

    bullet_style = ParagraphStyle(
        "Bullet",
        parent=body_style,
        leftIndent=18,
        bulletIndent=6,
        spaceAfter=2,
        alignment=TA_LEFT,
    )

    end_marker_style = ParagraphStyle(
        "EndMarker",
        parent=styles["Normal"],
        fontName="Helvetica-Oblique",
        fontSize=9,
        leading=12,
        alignment=TA_CENTER,
        textColor=HexColor("#888888"),
        spaceBefore=8,
        spaceAfter=8,
    )

    section_divider_style = ParagraphStyle(
        "SectionDivider",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=14,
        leading=20,
        alignment=TA_CENTER,
        textColor=HexColor("#1A5276"),
        spaceBefore=16,
        spaceAfter=16,
    )

    ascii_style = ParagraphStyle(
        "Ascii",
        parent=styles["Normal"],
        fontName="Courier",
        fontSize=7,
        leading=8,
        leftIndent=12,
        rightIndent=12,
        spaceBefore=4,
        spaceAfter=4,
        textColor=HexColor("#333333"),
    )

    toc_h1 = ParagraphStyle(
        "TOCHeading1",
        fontName="Helvetica-Bold",
        fontSize=11,
        leading=16,
        leftIndent=0,
        textColor=HexColor("#1A5276"),
    )
    toc_h2 = ParagraphStyle(
        "TOCHeading2",
        fontName="Helvetica",
        fontSize=10,
        leading=14,
        leftIndent=16,
        textColor=HexColor("#333333"),
    )

    # Custom doc template to support TOC entries + bookmarks.
    # Only H1 (Parts) and H2 (Articles/Introductions) are added to the bookmark
    # outline tree and TOC. H3 entries are bookmarked (for cross-reference) but
    # NOT added to the outline — this avoids "can't jump outline level" errors
    # when an H3 appears before any H2 (e.g. Preamble's "Identity" section).
    class TocDocTemplate(BaseDocTemplate):
        def afterFlowable(self, flowable):
            if isinstance(flowable, Paragraph):
                style_name = flowable.style.name
                text = flowable.getPlainText()
                if style_name == "H1":
                    key = f"h1-{self.seq.nextf('toc_h1')}"
                    self.canv.bookmarkPage(key)
                    self.canv.addOutlineEntry(text, key, level=0, closed=False)
                    self.notify("TOCEntry", (0, text, self.page, key))
                elif style_name == "H2":
                    key = f"h2-{self.seq.nextf('toc_h2')}"
                    self.canv.bookmarkPage(key)
                    self.canv.addOutlineEntry(text, key, level=1, closed=False)
                    self.notify("TOCEntry", (1, text, self.page, key))
                elif style_name == "H3":
                    key = f"h3-{self.seq.nextf('toc_h3')}"
                    self.canv.bookmarkPage(key)

    # Canvas subclass to draw page numbers in footer (two-pass: total pages)
    from reportlab.pdfgen import canvas as canvas_module

    class NumberedCanvas(canvas_module.Canvas):
        def __init__(self, *args, **kwargs):
            super().__init__(*args, **kwargs)
            self._saved_states = []

        def showPage(self):
            self._saved_states.append(dict(self.__dict__))
            super().showPage()

        def save(self):
            n = len(self._saved_states)
            for state in self._saved_states:
                self.__dict__.update(state)
                self._draw_footer(n)
                super().showPage()
            super().save()

        def _draw_footer(self, total):
            self.saveState()
            self.setStrokeColor(HexColor("#CCCCCC"))
            self.setLineWidth(0.5)
            self.line(0.75 * inch, 0.65 * inch, PAGE_W - 0.75 * inch, 0.65 * inch)
            self.setFont("Helvetica", 8)
            self.setFillColor(HexColor("#888888"))
            self.drawString(
                0.75 * inch,
                0.45 * inch,
                f"MITHQAL Blueprint {VERSION} \u2014 Constitutional Monetary Infrastructure Specification",
            )
            self.drawRightString(
                PAGE_W - 0.75 * inch,
                0.45 * inch,
                f"Page {self._pageNumber} of {total}",
            )
            self.restoreState()

    # Build story
    story = []

    # Cover page
    story.append(Spacer(1, 2.5 * inch))
    story.append(Paragraph(DOC_TITLE, title_style))
    story.append(Paragraph(VERSION, subtitle_style))
    story.append(Spacer(1, 0.3 * inch))
    story.append(Paragraph(DOC_SUBTITLE, subtitle_style))
    story.append(Spacer(1, 1.0 * inch))
    story.append(Paragraph(
        f"Version: {VERSION} (v18 base + v19 constitutional evolution)<br/>"
        f"Date: {DATE_STR}<br/>"
        f"5 Parts (Layers) &middot; 49 Articles &middot; 28,456 source lines",
        cover_meta_style,
    ))
    story.append(PageBreak())

    # Table of Contents page
    story.append(Paragraph("Table of Contents", h1_style))
    story.append(Spacer(1, 0.2 * inch))

    toc = TableOfContents()
    toc.levelStyles = [toc_h1, toc_h2]
    story.append(toc)
    story.append(PageBreak())

    def esc_pdf(s):
        return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    in_ascii = False
    ascii_buffer = []

    def flush_ascii():
        nonlocal in_ascii, ascii_buffer
        if in_ascii and ascii_buffer:
            txt = "<br/>".join(esc_pdf(l) for l in ascii_buffer)
            story.append(Paragraph(txt, ascii_style))
            ascii_buffer = []
            in_ascii = False

    for typ, content in parsed:
        if typ == "title":
            continue
        if typ == "blank":
            if in_ascii:
                ascii_buffer.append("")
            continue
        if typ == "ascii_art":
            if not in_ascii:
                in_ascii = True
                ascii_buffer = []
            ascii_buffer.append(content)
            continue
        else:
            flush_ascii()

        if typ == "part":
            story.append(Paragraph(esc_pdf(content), h1_style))
        elif typ in ("introduction", "article"):
            story.append(Paragraph(esc_pdf(content), h2_style))
        elif typ in ("invariant", "subheader", "numbered"):
            story.append(Paragraph(esc_pdf(content), h3_style))
        elif typ == "section_divider":
            story.append(Paragraph(esc_pdf(content), section_divider_style))
        elif typ == "end_marker":
            story.append(Paragraph(esc_pdf(content), end_marker_style))
        elif typ == "bullet":
            text = esc_pdf(content)
            story.append(Paragraph(text, bullet_style, bulletText="\u2022"))
        elif typ == "paragraph":
            story.append(Paragraph(esc_pdf(content), body_style))

    flush_ascii()

    # Build the document
    frame = Frame(
        0.85 * inch,
        0.85 * inch,
        PAGE_W - 1.7 * inch,
        PAGE_H - 1.7 * inch,
        id="main",
        leftPadding=0,
        rightPadding=0,
        topPadding=0,
        bottomPadding=0,
    )

    template = PageTemplate(id="main", frames=[frame])
    doc = TocDocTemplate(
        pdf_path,
        pagesize=letter,
        pageTemplates=[template],
        title=f"{DOC_TITLE} {VERSION}",
        author="MITHQAL Constitutional Council",
        subject=DOC_SUBTITLE,
    )

    doc.multiBuild(story, canvasmaker=NumberedCanvas)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    os.makedirs(OUT_DIR, exist_ok=True)

    print(f"Reading {BLUEPRINT_PATH} ...", flush=True)
    parsed = parse_blueprint(BLUEPRINT_PATH)
    total_elements = len(parsed)
    print(f"  Parsed {total_elements} elements", flush=True)

    counts = Counter(t for t, _ in parsed)
    print(f"  Element breakdown: {dict(counts)}", flush=True)

    print(f"\n[1/4] Generating Markdown -> {MD_PATH}", flush=True)
    md_lines = generate_markdown(parsed, MD_PATH)
    md_size = os.path.getsize(MD_PATH)
    print(f"  Wrote {md_lines} lines, {md_size:,} bytes", flush=True)

    print(f"\n[2/4] Generating HTML -> {HTML_PATH}", flush=True)
    html_lines = generate_html(parsed, HTML_PATH)
    html_size = os.path.getsize(HTML_PATH)
    print(f"  Wrote {html_lines} lines, {html_size:,} bytes", flush=True)

    print(f"\n[3/4] Generating DOCX -> {DOCX_PATH}", flush=True)
    generate_docx(parsed, DOCX_PATH)
    docx_size = os.path.getsize(DOCX_PATH)
    print(f"  Wrote {docx_size:,} bytes", flush=True)

    print(f"\n[4/4] Generating PDF -> {PDF_PATH}", flush=True)
    generate_pdf(parsed, PDF_PATH)
    pdf_size = os.path.getsize(PDF_PATH)
    print(f"  Wrote {pdf_size:,} bytes", flush=True)

    try:
        import pypdf
        reader = pypdf.PdfReader(PDF_PATH)
        pages = len(reader.pages)
        print(f"  PDF page count: {pages}", flush=True)
    except Exception as e:
        print(f"  WARNING: Could not read PDF page count: {e}", flush=True)

    print("\n=== DONE ===", flush=True)
    print(f"  MD:   {MD_PATH} ({md_size:,} bytes)")
    print(f"  HTML: {HTML_PATH} ({html_size:,} bytes)")
    print(f"  DOCX: {DOCX_PATH} ({docx_size:,} bytes)")
    print(f"  PDF:  {PDF_PATH} ({pdf_size:,} bytes)")


if __name__ == "__main__":
    main()
