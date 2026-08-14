#!/usr/bin/env python3
"""
MITHQAL v25.0 — Generate .docx from the FINAL blueprint markdown.
Uses python-docx for reliable, well-formatted output.
"""
import re
import os
import sys
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError:
    print("Installing python-docx...")
    os.system(f"{sys.executable} -m pip install python-docx --quiet")
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

INPUT = "/home/z/my-project/docs/blueprint/mithqal-v25-FINAL-blueprint.md"
OUTPUT = "/home/z/my-project/docs/blueprint/mithqal-v25-FINAL-blueprint.docx"

# Color palette (institutional: dark navy + gold)
COLOR_PRIMARY = RGBColor(0x1A, 0x36, 0x5F)      # Deep navy
COLOR_GOLD = RGBColor(0xC5, 0x9B, 0x1E)          # Institutional gold
COLOR_DARK = RGBColor(0x2D, 0x2D, 0x2D)          # Near-black
COLOR_GRAY = RGBColor(0x6B, 0x6B, 0x6B)          # Medium gray
COLOR_ACCENT = RGBColor(0x00, 0x6B, 0x3C)        # Emerald (for PASS)
COLOR_FAIL = RGBColor(0xB0, 0x1E, 0x1E)          # Crimson (for FAIL)
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_horizontal_rule(doc, color="C59B1E", size="12"):
    """Add a horizontal rule (bottom border on a paragraph)."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="{size}" w:space="1" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p

def parse_markdown_table(lines):
    """Parse a markdown table block into rows of cells."""
    rows = []
    for line in lines:
        line = line.strip()
        if not line.startswith("|"):
            continue
        # Split by |, remove first/last empty
        cells = [c.strip() for c in line.split("|")[1:-1]]
        rows.append(cells)
    # Remove separator row (contains --- or ===)
    rows = [r for r in rows if not all(re.match(r'^[-:]+$', c) for c in r)]
    return rows

def add_styled_table(doc, rows, header=True):
    """Add a formatted table from parsed rows."""
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j >= n_cols:
                continue
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(cell_text)
            run.font.size = Pt(9)
            run.font.name = "Calibri"
            if header and i == 0:
                run.font.bold = True
                run.font.color.rgb = COLOR_WHITE
                set_cell_shading(cell, "1A365F")
            else:
                run.font.color.rgb = COLOR_DARK
                if i % 2 == 0:
                    set_cell_shading(cell, "F0F0F0")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def main():
    print(f"Reading: {INPUT}")
    with open(INPUT, "r", encoding="utf-8") as f:
        content = f.read()

    doc = Document()

    # Page setup (A4)
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10.5)
    font.color.rgb = COLOR_DARK
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.3
    pf.space_after = Pt(4)

    # Heading styles
    for level, (size, color) in enumerate([(18, COLOR_PRIMARY), (14, COLOR_PRIMARY), (12, COLOR_PRIMARY), (11, COLOR_DARK)], 1):
        h_style = doc.styles[f'Heading {level}']
        h_style.font.name = 'Calibri'
        h_style.font.size = Pt(size)
        h_style.font.bold = True
        h_style.font.color.rgb = color
        h_style.paragraph_format.space_before = Pt(12)
        h_style.paragraph_format.space_after = Pt(6)
        h_style.paragraph_format.keep_with_next = True

    # Title style
    title_style = doc.styles['Title']
    title_style.font.name = 'Calibri'
    title_style.font.size = Pt(28)
    title_style.font.bold = True
    title_style.font.color.rgb = COLOR_PRIMARY

    lines = content.split("\n")
    i = 0
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Code blocks
        if stripped.startswith("```"):
            if in_code_block:
                # End code block
                code_text = "\n".join(code_lines)
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                run = p.add_run(code_text)
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                run.font.color.rgb = COLOR_GRAY
                # Light gray background
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
                p._p.get_or_add_pPr().append(shading)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # HTML comments — skip
        if stripped.startswith("<!--") or stripped.endswith("-->"):
            i += 1
            continue

        # Horizontal rules (--- or ===)
        if stripped in ("---", "===") or (stripped and all(c == "-" for c in stripped) and len(stripped) > 3):
            add_horizontal_rule(doc)
            i += 1
            continue

        # Tables
        if stripped.startswith("|") and i + 1 < len(lines) and "|---" in lines[i + 1]:
            # Collect table lines
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = parse_markdown_table(table_lines)
            add_styled_table(doc, rows)
            doc.add_paragraph()  # spacing after table
            continue

        # H1 (# )
        if stripped.startswith("# ") and not stripped.startswith("## "):
            text = stripped[2:].strip()
            # Check if it's a major part header
            if "PART" in text or "APPEND" in text or "COVER" in text or "INDEX" in text:
                doc.add_page_break()
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text)
                run.font.size = Pt(22)
                run.font.bold = True
                run.font.color.rgb = COLOR_PRIMARY
                add_horizontal_rule(doc, "1A365F", "18")
            else:
                doc.add_heading(text, level=1)
            i += 1
            continue

        # H2 (## )
        if stripped.startswith("## ") and not stripped.startswith("### "):
            text = stripped[3:].strip()
            doc.add_heading(text, level=2)
            i += 1
            continue

        # H3 (### )
        if stripped.startswith("### ") and not stripped.startswith("#### "):
            text = stripped[4:].strip()
            doc.add_heading(text, level=3)
            i += 1
            continue

        # H4 (#### )
        if stripped.startswith("#### "):
            text = stripped[5:].strip()
            doc.add_heading(text, level=4)
            i += 1
            continue

        # Blockquotes (> )
        if stripped.startswith("> "):
            text = stripped[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.right_indent = Cm(1)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(text)
            run.font.size = Pt(11)
            run.font.italic = True
            run.font.color.rgb = COLOR_PRIMARY
            run.font.bold = True
            # Left border
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:left w:val="single" w:sz="24" w:space="8" w:color="C59B1E"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
            # Light gold background
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="FFFBF0"/>')
            pPr.append(shading)
            i += 1
            continue

        # Bullet lists (- or •)
        if stripped.startswith("- ") or stripped.startswith("• "):
            text = stripped[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            # Handle bold markers **text**
            add_formatted_runs(p, text)
            p.paragraph_format.left_indent = Cm(0.75)
            i += 1
            continue

        # Numbered lists (1. 2. etc)
        if re.match(r'^\d+\.\s', stripped):
            text = re.sub(r'^\d+\.\s', '', stripped)
            p = doc.add_paragraph(style='List Number')
            add_formatted_runs(p, text)
            p.paragraph_format.left_indent = Cm(0.75)
            i += 1
            continue

        # Empty lines
        if not stripped:
            i += 1
            continue

        # Regular paragraph with bold/italic formatting
        p = doc.add_paragraph()
        add_formatted_runs(p, stripped)
        i += 1

    print(f"Writing: {OUTPUT}")
    doc.save(OUTPUT)
    size = os.path.getsize(OUTPUT)
    print(f"Done. Size: {size:,} bytes ({size/1024:.1f} KB)")

def add_formatted_runs(paragraph, text):
    """Add runs with **bold** and *italic* formatting."""
    # Split by ** for bold
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            # Bold
            inner = part[2:-2]
            # Check for italic within bold
            sub_parts = re.split(r'(\*.*?\*)', inner)
            for sp in sub_parts:
                if sp.startswith("*") and sp.endswith("*") and len(sp) > 2:
                    run = paragraph.add_run(sp[1:-1])
                    run.font.bold = True
                    run.font.italic = True
                else:
                    run = paragraph.add_run(sp)
                    run.font.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            # Italic
            run = paragraph.add_run(part[1:-1])
            run.font.italic = True
        else:
            # Check for inline code `text`
            code_parts = re.split(r'(`.*?`)', part)
            for cp in code_parts:
                if cp.startswith("`") and cp.endswith("`"):
                    run = paragraph.add_run(cp[1:-1])
                    run.font.name = "Consolas"
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = COLOR_GRAY
                else:
                    if cp:
                        run = paragraph.add_run(cp)

if __name__ == "__main__":
    main()
